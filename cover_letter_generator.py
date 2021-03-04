from notion.client import NotionClient
import notion_credentials
from docx import Document
from docx2pdf import convert

#Takes string parameters to fill out file
#Build Microsoft Word file using python-docx 
#Save as DOCX and PDF using docx2pdf 
def _buildCoverLetter(company, role, where_you_found_job, reason_you_want_job):
    document = Document()
    
    cover_letter = f"Rhett Reisman {company} {role}"
    
    p_explanation = document.add_paragraph('This cover letter has been produced by my Notion Cover Letter Generator. ')
    p_systems = document.add_paragraph('I developed this script because I believe that reusable systems are a lot more powerful than individual goals.')
    p_longterm = document.add_paragraph('I’m hoping that by including this preface I can find a company that prioritizes systems thinking and long-term strategy over short-term goals.')
    p_source_code = document.add_paragraph('Source code is available here: https://github.com/rhettre/NotionCoverLetterGenerator')

    p_divider = document.add_paragraph('______________________________________________________________________________')

    p_hello = document.add_paragraph(f'Hello {company} Team,')
    p_where_you_saw_job = document.add_paragraph(f'I saw your job posting on {where_you_found_job} and it sounds like an interesting opportunity!')
    p_background = document.add_paragraph(f'A little background on me: I grew up loving math and technology which led me to a degree in Computer Science and Business from Lehigh University. This led to my first job in insurance consulting where I have worked on multi-million dollar software implementations first as a developer and now as a project manager.')
    p_why = document.add_paragraph(f'Throughout my career I’ve taken on more and more responsibility and I see this {role} role at {company} as a great next step and huge opportunity. {reason_you_want_job}')
    p_any_questions = document.add_paragraph(f'Please let me know if I can answer any questions for you, I’d love to learn more about the position and your work at {company}')
    p_sign_off = document.add_paragraph('Cheers,')
    p_signature = document.add_paragraph('Rhett Reisman')
    document.save(f'{output_path}/{cover_letter}.docx')
    convert(f"{output_path}/{cover_letter}.docx", f"{output_path}/{cover_letter}.pdf")

client = NotionClient(token_v2 = notion_credentials.NOTION_TOKEN)
output_path = "/Users/rhettre/My Drive/Business/Resumes/Cover Letters"

#lookup ids for different columns
cover_letter_id = 'A_u<'
position_id = 'xyVl'
status_id = 'RpsQ'
location_id = 'SXHt'
resume_id = 'XW]^'
where_you_found_job_id = 'Hv_n'
why_this_job_id = 'WB`~'
company_id = 'title'

cv = client.get_collection_view(notion_credentials.NOTION_URL)

rows = cv.collection.get_rows()

for row in rows:
    #Collect all where status is ready for processing - complete information 
    if row.get_property(status_id) == "Information Complete":
        print(row.get_property(position_id))
        _buildCoverLetter(row.get_property(company_id), row.get_property(position_id), row.get_property(where_you_found_job_id), row.get_property(why_this_job_id))
